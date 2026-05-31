import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_sang_kien():
    doc = Document()
    
    # Title
    title = doc.add_heading('BÁO CÁO SÁNG KIẾN KINH NGHIỆM', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_heading('Tên sáng kiến: Giải pháp số hóa phễu bán hàng qua mô hình Kanban nhằm nâng cao tỷ lệ chuyển đổi khách hàng trên hệ thống BCRM', level=1)
    
    doc.add_paragraph('Tác giả: [Tên của bạn]\nĐơn vị: [Tên chi nhánh/phòng ban]\nNgày lập: [Ngày/Tháng/Năm]')
    
    doc.add_heading('1. Thực trạng trước khi áp dụng sáng kiến', level=2)
    p1 = doc.add_paragraph()
    p1.add_run('Trước khi hệ thống BCRM được đưa vào sử dụng, công tác quản lý khách hàng và theo dõi bán hàng tại đơn vị gặp một số khó khăn:\n').bold = True
    doc.add_paragraph('- Quản lý phân tán: Dữ liệu khách hàng được quản lý thủ công qua sổ sách hoặc các file Excel cá nhân, dễ dẫn đến mất mát dữ liệu hoặc trùng lặp khi có sự thay đổi nhân sự.', style='List Bullet')
    doc.add_paragraph('- Thiếu bức tranh tổng thể: Khó khăn trong việc theo dõi tiến độ của từng cơ hội bán hàng (Deal). Việc không biết chính xác khách hàng đang ở giai đoạn nào (Tiếp cận, Đàm phán, hay Đã chốt) làm giảm hiệu quả bám sát và chăm sóc.', style='List Bullet')
    doc.add_paragraph('- Báo cáo thủ công: Mất nhiều thời gian để tổng hợp KPI cuối tuần/cuối tháng. Quản lý khó nắm bắt tình hình tức thời để đốc thúc nhân viên.', style='List Bullet')
    
    doc.add_heading('2. Nội dung chi tiết của sáng kiến', level=2)
    p2 = doc.add_paragraph()
    p2.add_run('Sáng kiến đề xuất số hóa toàn diện quy trình bán hàng bằng hệ thống BCRM, với trọng tâm là ứng dụng mô hình Kanban vào quản lý cơ hội kinh doanh:\n').bold = True
    doc.add_paragraph('- Trực quan hóa phễu bán hàng (Kanban Board): Toàn bộ vòng đời của một quy trình sales được chia thành các cột trạng thái rõ ràng (Tiếp cận -> Đàm phán -> Ký hợp đồng -> Hoàn thành).', style='List Bullet')
    doc.add_paragraph('- Quản lý bằng thao tác kéo-thả (Drag & Drop): Chuyên viên dễ dàng cập nhật tiến độ của từng khách hàng chỉ bằng một thao tác kéo thả, giảm thiểu tối đa thời gian nhập liệu truyền thống.', style='List Bullet')
    doc.add_paragraph('- Dashboard theo dõi KPI Real-time: Hệ thống tự động tổng hợp dư nợ cho vay, tổng vốn huy động và số lượng lịch hẹn cần xử lý hiển thị ngay trên màn hình chính.', style='List Bullet')

    doc.add_heading('3. Hiệu quả mang lại', level=2)
    doc.add_paragraph('- Đối với chuyên viên: Tiết kiệm 30% thời gian làm báo cáo hằng ngày. Tăng tỷ lệ chuyển đổi (Win-rate) do không bỏ sót các lịch hẹn và khách hàng đang ở giai đoạn "nóng".', style='List Bullet')
    doc.add_paragraph('- Đối với cấp quản lý: Quản lý minh bạch, đánh giá chính xác năng lực của từng nhân sự qua các chỉ số tương tác và KPI trực tuyến.', style='List Bullet')
    doc.add_paragraph('- Đối với tính bảo mật: Toàn bộ thông tin khách hàng được lưu trữ tập trung, phân quyền rõ ràng, tránh rủi ro rò rỉ dữ liệu.', style='List Bullet')
    
    doc.add_heading('4. Phạm vi và đối tượng áp dụng', level=2)
    doc.add_paragraph('Sáng kiến đã được áp dụng hiệu quả cho toàn bộ chuyên viên kinh doanh và cấp quản lý trực tiếp tại đơn vị.')
    
    os.makedirs('docs', exist_ok=True)
    doc.save('docs/Sang_Kien_Kinh_Nghiem_BCRM_v2.docx')

def create_hdsd():
    doc = Document()
    
    title = doc.add_heading('HƯỚNG DẪN SỬ DỤNG BCRM - LUỒNG CÔNG VIỆC & PHÂN QUYỀN', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph('Tài liệu hướng dẫn sử dụng hệ thống BCRM cho các cấp bậc: Cấp quản trị (Admin Level 1, Level 2) và Chuyên viên kinh doanh (User).')
    
    doc.add_heading('PHẦN 1: DÀNH CHO CẤP QUẢN LÝ (ADMIN LEVEL 1 & LEVEL 2)', level=1)
    
    doc.add_heading('1. Đầu kỳ (Tháng/Tuần): Giao và Phân bổ Kế hoạch', level=2)
    doc.add_paragraph('Với Admin Level 1 (Giám đốc/Ban Giám đốc):').bold = True
    doc.add_paragraph('- Bước 1: Vào mục "Kế hoạch". Thực hiện thiết lập và giao chỉ tiêu tổng (Kế hoạch tháng/quý) xuống các phòng ban (Admin Level 2).', style='List Bullet')
    doc.add_paragraph('- Bước 2: Theo dõi trạng thái tiếp nhận kế hoạch của các phòng ban.', style='List Bullet')
    
    doc.add_paragraph('Với Admin Level 2 (Trưởng phòng/Trưởng nhóm):').bold = True
    doc.add_paragraph('- Bước 1: Xem chỉ tiêu phòng ban được giao từ Level 1.', style='List Bullet')
    doc.add_paragraph('- Bước 2: Dựa trên năng lực nhân sự, phân bổ kế hoạch cụ thể (chỉ tiêu huy động, cho vay, khách hàng mới) xuống từng Chuyên viên trong phòng.', style='List Bullet')
    doc.add_paragraph('- Bước 3: Có thể tạo và giao thêm các Kế hoạch ngày/tuần chi tiết để nhân viên thực hiện.', style='List Bullet')

    doc.add_heading('2. Quản trị hàng ngày: Theo dõi tiến độ', level=2)
    doc.add_paragraph('- Bước 1: Quản lý truy cập Dashboard để xem bức tranh tổng thể: Khách hàng quản lý, Dư nợ, Vốn huy động của toàn bộ nhân viên thuộc quyền.', style='List Bullet')
    doc.add_paragraph('- Bước 2: Mục "Đội ngũ nhân viên": Quan sát hiệu suất của từng chuyên viên (Số lượng KH, Số tương tác, Kết quả bán hàng).', style='List Bullet')
    doc.add_paragraph('- Bước 3: Xem báo cáo phân bổ và tiến độ thực hiện kế hoạch (Tỷ lệ % hoàn thành KPIs) để đưa ra đánh giá năng lực.', style='List Bullet')

    doc.add_heading('3. Điều hành và Đốc thúc', level=2)
    doc.add_paragraph('- Từ kết quả báo cáo trên hệ thống, thực hiện nhắc nhở các chuyên viên có tiến độ chạy KPI chậm.', style='List Bullet')
    doc.add_paragraph('- Quản lý có thể xem chi tiết Bảng Kanban bán hàng của nhân viên để hỗ trợ giải quyết các khách hàng khó/kẹt ở bước đàm phán.', style='List Bullet')


    doc.add_heading('PHẦN 2: DÀNH CHO CHUYÊN VIÊN KINH DOANH (USER)', level=1)
    
    doc.add_heading('1. Bắt đầu ngày mới: Nhận Kế hoạch & Mục tiêu', level=2)
    doc.add_paragraph('- Bước 1: Đăng nhập vào hệ thống BCRM, truy cập module "Kế hoạch".')
    doc.add_paragraph('- Bước 2: Xem chi tiết các Kế hoạch tháng/tuần/ngày đã được Admin phân bổ cho mình.')
    doc.add_paragraph('- Bước 3: Ở Trang chủ (Dashboard), cuộn xuống phần "Lịch hẹn cần xử lý". Ưu tiên xử lý các lịch hẹn trong ngày hôm nay.')
    
    doc.add_heading('2. Trong ngày làm việc: Quản lý Khách hàng & Kế hoạch', level=2)
    doc.add_paragraph('Thực hiện nghiệp vụ Khách hàng:').bold = True
    doc.add_paragraph('- Khách hàng mới: Vào tab "Khách hàng" -> Nhấn "Thêm mới" -> Cập nhật thông tin.')
    doc.add_paragraph('- Ghi nhận cơ hội: Gắn khách hàng với nhu cầu cụ thể (Tiền gửi/Khoản vay).')
    
    doc.add_paragraph('Quản lý tiến độ bán hàng qua Kanban:').bold = True
    doc.add_paragraph('- Vào mục Kanban (Sales Support).')
    doc.add_paragraph('- Kéo thả trạng thái Deal (Tiếp cận -> Đàm phán -> Hoàn thành) tuỳ theo tình hình chốt sales.')
    doc.add_paragraph('- Cập nhật "Tương tác": Ghi chú lại nội dung gặp/gọi điện thoại với khách.')

    doc.add_paragraph('Cập nhật tiến độ kế hoạch được giao:').bold = True
    doc.add_paragraph('- Liên tục kiểm tra tiến độ của bản thân so với kế hoạch quản lý đã giao.')
    
    doc.add_heading('3. Cuối ngày: Đánh giá cá nhân', level=2)
    doc.add_paragraph('- Truy cập "KPI Summary Table" (Bảng tổng hợp KPI) để kiểm tra tổng Dư nợ và Vốn huy động trong ngày/tháng.', style='List Bullet')
    doc.add_paragraph('- Tự đánh giá khả năng hoàn thành chỉ tiêu được giao và chuẩn bị công việc cho ngày hôm sau.', style='List Bullet')
    
    doc.add_paragraph('\nLưu ý: Hệ thống được liên thông từ cấp lãnh đạo đến chuyên viên. Các nỗ lực tương tác và kết quả kinh doanh của bạn sẽ được ghi nhận và hiển thị ngay lập tức tới cấp Quản lý!', style='Intense Quote')
    
    os.makedirs('docs', exist_ok=True)
    doc.save('docs/HDSD_BCRM_Workflow_v2.docx')

if __name__ == "__main__":
    create_sang_kien()
    create_hdsd()
    print("Successfully generated docs in the 'docs' folder.")
